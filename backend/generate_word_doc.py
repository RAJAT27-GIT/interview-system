import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_massive_technical_report():
    doc = Document()

    # --- GLOBAL STYLE ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    def add_page_title(text):
        doc.add_page_break()
        h = doc.add_heading(text, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('\n')

    def add_section_title(text):
        h = doc.add_heading(text, level=2)
        doc.add_paragraph('\n')

    def add_para(text):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    # --- FRONT PAGE ---
    for _ in range(3): doc.add_paragraph()
    t = doc.add_heading('A TECHNICAL PROJECT REPORT ON', level=2)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = doc.add_heading('AI-POWERED MULTI-ROUND INTERVIEW SYSTEM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n' * 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENTS\nFOR THE AWARD OF THE DEGREE OF')
    run.bold = True
    doc.add_paragraph('\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BACHELOR OF TECHNOLOGY\nIN\nCOMPUTER SCIENCE & ENGINEERING')
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph('\n' * 3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SUBMITTED BY:\n[Your Name]\nRoll No: [Your Roll No]')
    run.italic = True
    doc.add_paragraph('\n' * 4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('UNDER THE GUIDANCE OF:\n[Supervisor Name]\n[Designation]')
    run.bold = True
    doc.add_paragraph('\n' * 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING\n[YOUR INSTITUTE NAME]')
    run.bold = True

    # --- CERTIFICATE, ACK, ABSTRACT ---
    add_page_title('CERTIFICATE')
    add_para("This is to certify that the project work entitled 'AI-Powered Multi-Round Interview System' "
             "is a bonafide work carried out by [Your Name] in partial fulfillment of the requirements "
             "for the award of degree of Bachelor of Technology in Computer Science & Engineering.")
    doc.add_paragraph('\n' * 10)
    doc.add_paragraph('Signature of HOD\t\t\t\tSignature of Supervisor')

    add_page_title('ACKNOWLEDGEMENT')
    add_para("I would like to express my deepest appreciation to all those who provided me the "
             "possibility to complete this report. A special thanks to my supervisor, [Supervisor Name], "
             "for their time and valuable suggestions. I also thank my parents and friends for their "
             "continuous support.")

    add_page_title('ABSTRACT')
    add_para("In the era of rapid technological advancement, the recruitment industry is witnessing "
             "a shift towards automation and intelligence. This project presents an 'AI-Powered Multi-Round "
             "Interview System' that leverages Large Language Models (LLMs) and advanced speech recognition "
             "to automate the preliminary technical screening process. By parsing resumes, generating "
             "dynamic questions, and providing live code evaluation, the system offers a standardized "
             "and objective assessment tool for organizations.")

    # --- PHASE 3: DEEP DIVE INTO TECHNOLOGIES ---
    add_page_title('TECHNOLOGY DEEP DIVE: FRONTEND')
    
    add_section_title('1. React.js (v18+)')
    add_para("React is the primary frontend library used for building the user interface. It is chosen "
             "for its declarative nature and component-based architecture. In this project, React "
             "manages the complex state of the interview lifecycle, including the transitions between "
             "theory, coding, and behavioral rounds. The use of Functional Components and Hooks (useState, "
             "useEffect, useContext) ensures that the application remains modular and easy to debug.")

    add_section_title('2. Vite Build Tool')
    add_para("Vite is used as the frontend build tool and development server. Unlike traditional tools "
             "like Create React App (CRA), Vite leverages ES modules to provide near-instant hot module "
             "replacement (HMR). This significantly improved the development speed for complex UI elements "
             "like the live coding editor and audio visualization components.")

    add_section_title('3. Tailwind CSS & Framer Motion')
    add_para("For styling, Tailwind CSS is used to implement a utility-first design system. This "
             "ensures that the UI is responsive across all devices. Framer Motion is integrated "
             "to provide fluid animations, such as the smooth transitions between interview questions "
             "and the pulsing animation of the audio recording button, enhancing the overall user experience.")

    add_page_title('TECHNOLOGY DEEP DIVE: BACKEND')
    
    add_section_title('4. FastAPI (Python Framework)')
    add_para("FastAPI is the backbone of our backend services. It is a modern, fast (high-performance) "
             "web framework for building APIs with Python. We utilized FastAPI's Pydantic integration "
             "for robust data validation and its asynchronous capabilities to handle long-running "
             "AI generation tasks without blocking the server's main event loop.")

    add_section_title('5. Groq Cloud & Llama 3.3')
    add_para("The core intelligence of the system is provided by Meta's Llama 3.3-70B model, "
             "running on the Groq Cloud platform. Groq's Language Processing Unit (LPU) technology "
             "allows us to achieve inference speeds that were previously impossible, ensuring that "
             "the AI's 'thinking time' is virtually imperceptible to the candidate.")

    add_section_title('6. Whisper V3 & gTTS')
    add_para("For audio interaction, we integrated OpenAI's Whisper V3 (via Groq) for transcription. "
             "It converts the candidate's spoken words into text with high accuracy, even in noisy "
             "environments. Conversely, gTTS (Google Text-to-Speech) is used to read out questions "
             "to the candidate, providing a multi-modal interview experience.")

    add_page_title('TECHNOLOGY DEEP DIVE: DATABASE & UTILITIES')

    add_section_title('7. MongoDB Atlas')
    add_para("MongoDB is our choice for persistent data storage. As a NoSQL database, it handles the "
             "dynamic nature of interview sessions—where each round might have different metadata—very "
             "effectively. We use it to store user profiles, generated question banks, and activity logs "
             "for administrative review.")

    add_section_title('8. SpaCy & Pdfplumber')
    add_para("The resume parsing logic relies on SpaCy, an industrial-strength NLP library, and "
             "pdfplumber. While pdfplumber handles the structural extraction of text from PDF files, "
             "SpaCy performs Named Entity Recognition (NER) to identify professional skills, education "
             "history, and contact information automatically.")

    # --- PHASE 4: COMPONENT & PAGE ARCHITECTURE ---
    add_page_title('DETAILED PAGE & COMPONENT ARCHITECTURE')

    add_section_title('1. Registration & Auth Pages')
    add_para("The entry point for candidates. It collects basic profile information and stores it "
             "securely. The authentication logic is handled via JWT (JSON Web Tokens), ensuring "
             "that only authorized users can access the interview sessions.")

    add_section_title('2. Candidate Dashboard')
    add_para("A central hub where users can see their previous performance, upload their latest resume, "
             "and select the difficulty tier for their next interview. It uses data visualization "
             "libraries to show score trends and skill breakdowns.")

    add_section_title('3. The Interview Room')
    add_para("The most critical part of the UI. It includes a dynamic question display, an audio "
             "controller for voice responses, and a integrated Monaco Editor for coding tasks. "
             "The room is synchronized with a backend timer that enforces strict round durations.")

    add_section_title('4. Proctoring Modal & Alerts')
    add_para("A specialized component that remains active throughout the session. It monitors "
             "the 'Visibility State' of the browser tab. If a user switches tabs, it triggers an "
             "immediate warning overlay, preventing the user from interacting with the interview "
             "until they acknowledge the violation.")

    # --- PHASE 5: LOGIC & ALGORITHMS ---
    add_page_title('ALGORITHMS & LOGIC DEEP DIVE')

    add_section_title('1. Resume Parsing Pipeline')
    add_para("Step 1: Raw Text Extraction from PDF.\n"
             "Step 2: Cleaning (Removing special characters, stop-words).\n"
             "Step 3: Keyword Extraction using a predefined Skill-Tree.\n"
             "Step 4: ATS Scoring based on keyword density and structural completeness.")

    add_section_title('2. Dynamic Question Generation (LLM Logic)')
    add_para("The system uses a 'Multi-shot Prompting' technique. It sends the candidate's skills "
             "to the AI along with strict instructions to output valid JSON. The AI generates "
             "questions that specifically target the technologies listed in the resume, ensuring "
             "that the interview is relevant and personalized.")

    add_section_title('3. Automated Code Evaluation Sandbox')
    add_para("When a user submits code, the backend initiates a secure subprocess. It pipes the "
             "user's code into a virtual file, attaches the required test case inputs, and captures "
             "the standard output. If the output matches the 'Expected Output', the test case passes. "
             "The system handles multiple languages by selecting the appropriate runtime (Python/Node).")

    # --- PHASE 6: SECURITY & PROJECT MANAGEMENT ---
    add_page_title('SECURITY, RELIABILITY & FUTURE SCOPE')

    add_section_title('1. Security Measures')
    add_para("We implemented several security layers: 1. Input Sanitization to prevent prompt injection. "
             "2. Resource Limiting in the code executor to prevent CPU exhaustion. 3. Secure "
             "Environment Variables for API key protection.")

    add_section_title('2. Future Enhancements')
    add_para("The roadmap includes: 1. Real-time Video Monitoring for facial expression analysis. "
             "2. Integration with LinkedIn API for automatic profile importing. 3. Collaboration "
             "features for live human-AI hybrid interviewing.")

    add_page_title('CONCLUSION')
    add_para("This project successfully bridges the gap between traditional recruitment and AI. "
             "It offers a scalable solution for the modern workforce, ensuring that the best "
             "talent is identified through a fair, fast, and rigorous technical assessment process.")

    # Save the document
    file_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'AI_Interview_System_Massive_Technical_Report.docx')
    doc.save(file_path)
    print(f"File saved at: {file_path}")

if __name__ == "__main__":
    create_massive_technical_report()
