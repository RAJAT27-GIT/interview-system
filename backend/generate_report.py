import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize Document
doc = Document()

# Set standard styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_title_page():
    doc.add_spacer = lambda: doc.add_paragraph()
    
    for _ in range(5): doc.add_paragraph()
    
    title = doc.add_heading('REALISTIC AI INTERVIEW & RESUME SYSTEM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(3): doc.add_paragraph()
    
    p = doc.add_paragraph('A Project Report Submitted in Partial Fulfillment of the Requirements for the Award of')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('Bachelor of Technology')
    p.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('In')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('Computer Science & Engineering')
    p.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(10): doc.add_paragraph()
    
    p = doc.add_paragraph('Submitted By:')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('[Your Name Here]')
    p.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

def add_markdown_section(title, text):
    doc.add_heading(title, level=1)
    doc.add_paragraph(text)

def add_code_section(filename, filepath):
    doc.add_heading(f'File: {filename}', level=2)
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            code = f.read()
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    except Exception as e:
        doc.add_paragraph(f"Error reading file: {str(e)}")

# --- START GENERATION ---
add_title_page()

# Abstract
add_markdown_section('ABSTRACT', 'In the modern era of recruitment, traditional interview processes are often time-consuming and prone to human bias. This project presents an end-to-end solution for automated candidate evaluation using Generative AI...')

# Chapter 1
add_markdown_section('CHAPTER 1: INTRODUCTION', 'The recruitment landscape is evolving rapidly. Organizations are looking for smarter ways to filter candidates. Our project provides a "Human-in-the-loop" style AI that acts as a primary interviewer...')

# Chapter 2
add_markdown_section('CHAPTER 2: LITERATURE SURVEY', 'Existing systems often rely on static question banks which candidates can easily memorize. Our system uses LLMs to generate dynamic, resume-specific questions that test deep knowledge...')

# Chapter 3
add_markdown_section('CHAPTER 3: SYSTEM ANALYSIS', 'Software Requirements: Node.js, Python, FastAPI, React. Hardware Requirements: i5 Processor, 8GB RAM. Functional Requirements: Resume parsing, STT/TTS integration, Proctoring...')

# Chapter 5
add_markdown_section('CHAPTER 5: IMPLEMENTATION', 'The project is divided into two main parts: The Backend (FastAPI) and The Frontend (React). The backend handles AI logic, code execution, and data management. The frontend provides a premium user interface...')

# Appendix - Code (This is where most pages come from)
doc.add_page_break()
doc.add_heading('APPENDIX: SOURCE CODE', level=1)

base_path = r'c:\Users\asus\OneDrive\Desktop\MERN_SPI\project\interview-system'
files_to_include = [
    ('main.py', os.path.join(base_path, 'backend', 'main.py')),
    ('interview_engine.py', os.path.join(base_path, 'backend', 'app', 'interview_engine.py')),
    ('evaluator.py', os.path.join(base_path, 'backend', 'app', 'evaluator.py')),
    ('ResumeBuilder.jsx', os.path.join(base_path, 'frontend', 'src', 'pages', 'ResumeBuilder.jsx')),
    ('Interview.jsx', os.path.join(base_path, 'frontend', 'src', 'pages', 'Interview.jsx')),
    ('InterviewFlow.jsx', os.path.join(base_path, 'frontend', 'src', 'components', 'InterviewFlow.jsx')),
    ('CodeEditor.jsx', os.path.join(base_path, 'frontend', 'src', 'components', 'CodeEditor.jsx')),
]

for name, path in files_to_include:
    add_code_section(name, path)
    doc.add_page_break()

# Save
output_path = os.path.join(base_path, 'AI_Interview_System_Report.docx')
doc.save(output_path)
print(f"Report generated successfully at: {output_path}")
